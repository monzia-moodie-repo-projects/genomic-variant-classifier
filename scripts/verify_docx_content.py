#!/usr/bin/env python3
"""VERIFY_DOCX_CONTENT -- assert a generated .docx contains expected substrings in its
paragraph text. Avoids fragile PowerShell -c quote-escaping. Exit 0 = all present.

Usage: python verify_docx_content.py <docx_path> <needle1> [<needle2> ...]
"""
from __future__ import annotations
import sys


def main() -> int:
    if len(sys.argv) < 3:
        print("usage: python verify_docx_content.py <docx> <needle> [needle...]")
        return 2
    path, needles = sys.argv[1], sys.argv[2:]
    try:
        from docx import Document
    except ImportError:
        print("python-docx not installed"); return 2
    doc = Document(path)
    # include paragraphs AND table cells (python-docx splits some content into tables)
    parts = [p.text for p in doc.paragraphs]
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                parts.append(cell.text)
    text = "\n".join(parts)
    ok = True
    for n in needles:
        present = n in text
        print(f"  [{'PASS' if present else 'FAIL'}] {n}")
        ok &= present
    print("ALL PRESENT" if ok else "*** MISSING CONTENT ***")
    return 0 if ok else 1


if __name__ == "__main__":
    raise SystemExit(main())
