#!/usr/bin/env python3
"""
scripts/make_roadmap_docx.py
============================
Generate docs/ROADMAP.docx IN-PLACE from docs/ROADMAP.md (the source of
truth). This permanently removes the recurring "ROADMAP.docx not in
Downloads" install failure: the .docx is rebuilt locally from the committed
Markdown every session, so the two never drift.

Usage (from repo root):
    python scripts/make_roadmap_docx.py
    # or explicit paths:
    python scripts/make_roadmap_docx.py docs/ROADMAP.md docs/ROADMAP.docx

Strategy: prefer pandoc (best fidelity) if it is on PATH; otherwise fall back
to a self-contained python-docx converter (`pip install python-docx`).
"""
from __future__ import annotations

import re
import shutil
import subprocess
import sys
from pathlib import Path


def _via_pandoc(md: Path, docx: Path) -> bool:
    if shutil.which("pandoc") is None:
        return False
    try:
        subprocess.run(
            ["pandoc", str(md), "-f", "gfm", "-o", str(docx)],
            check=True, capture_output=True,
        )
        return True
    except subprocess.CalledProcessError as exc:
        sys.stderr.write(f"pandoc failed ({exc}); falling back to python-docx\n")
        return False


def _via_python_docx(md: Path, docx: Path) -> None:
    try:
        from docx import Document
    except ImportError:
        sys.exit(
            "python-docx not installed and pandoc not found.\n"
            "  pip install python-docx        (or install pandoc and re-run)"
        )

    doc = Document()
    bold = re.compile(r"\*\*(.+?)\*\*")
    code_inline = re.compile(r"`([^`]+)`")
    cell_sep = re.compile(r"^:?-{1,}:?$")

    def clean(s: str) -> str:
        return code_inline.sub(r"\1", bold.sub(r"\1", s))

    def split_cells(row: str) -> list[str]:
        return [c.strip() for c in row.strip().strip("|").split("|")]

    def is_sep(row: str) -> bool:
        parts = [c for c in split_cells(row) if c != ""]
        return bool(parts) and all(cell_sep.match(c) for c in parts)

    def add_table(header: list[str], rows: list[list[str]]) -> None:
        n = len(header)
        if n == 0:
            return
        t = doc.add_table(rows=1, cols=n)
        try:
            t.style = "Table Grid"        # built-in; always present in the default template
        except KeyError:
            pass
        for j, h in enumerate(header):
            t.rows[0].cells[j].paragraphs[0].add_run(clean(h)).bold = True
        for r in rows:
            rc = t.add_row().cells
            for j in range(n):
                rc[j].text = clean(r[j]) if j < len(r) else ""

    lines = md.read_text(encoding="utf-8").splitlines()
    i = 0
    in_code = False
    code_buf: list[str] = []

    def flush_code() -> None:
        if code_buf:
            p = doc.add_paragraph()
            p.add_run("\n".join(code_buf)).font.name = "Consolas"
            code_buf.clear()

    while i < len(lines):
        line = lines[i].rstrip("\n")
        st = line.strip()
        if st.startswith("```"):
            if in_code:
                flush_code()
            in_code = not in_code
            i += 1
            continue
        if in_code:
            code_buf.append(line)
            i += 1
            continue
        if not st:
            i += 1
            continue
        # GFM table: a leading-pipe header row immediately followed by a separator row
        if st.startswith("|") and i + 1 < len(lines) and is_sep(lines[i + 1]):
            header = split_cells(st)
            i += 2
            body: list[list[str]] = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                body.append(split_cells(lines[i].strip()))
                i += 1
            add_table(header, body)
            continue
        if st.startswith("### "):
            doc.add_heading(clean(st[4:].strip()), level=3)
        elif st.startswith("## "):
            doc.add_heading(clean(st[3:].strip()), level=2)
        elif st.startswith("# "):
            doc.add_heading(clean(st[2:].strip()), level=1)
        elif re.match(r"^[-*+]\s+", st):
            doc.add_paragraph(clean(st[2:].strip()), style="List Bullet")
        elif re.match(r"^\d+\.\s+", st):
            doc.add_paragraph(clean(re.sub(r"^\d+\.\s+", "", st)), style="List Number")
        else:
            doc.add_paragraph(clean(st))
        i += 1
    flush_code()
    doc.save(str(docx))


def main(argv: list[str]) -> None:
    md = Path(argv[1]) if len(argv) > 1 else Path("docs/ROADMAP.md")
    docx = Path(argv[2]) if len(argv) > 2 else Path("docs/ROADMAP.docx")

    if not md.exists():
        sys.exit(f"Source Markdown not found: {md} (run from the repo root).")

    docx.parent.mkdir(parents=True, exist_ok=True)

    if _via_pandoc(md, docx):
        print(f"Wrote {docx} via pandoc.")
    else:
        _via_python_docx(md, docx)
        print(f"Wrote {docx} via python-docx.")

    size = docx.stat().st_size
    print(f"  {docx}  ({size:,} bytes)")
    if size < 1024:
        sys.stderr.write("WARNING: output < 1 KB - verify ROADMAP.md had content.\n")


if __name__ == "__main__":
    main(sys.argv)
